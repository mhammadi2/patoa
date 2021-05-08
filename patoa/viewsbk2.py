from django.shortcuts import render, HttpResponse, redirect
import requests
from requests import get
from bs4 import BeautifulSoup
from django.http import HttpResponseRedirect
#from .models import Link
from .models import Patent
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver import Chrome
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
import os
import time
import io
from pathlib import Path
from webdriver_manager.chrome import ChromeDriverManager

from django.contrib import messages, auth
from django.contrib.auth.models import User

from django.contrib.auth.decorators import login_required

from .forms import PatentForm
from .forms import PatForm


# Build paths inside the project like this: BASE_DIR / 'subdir'.
BASE_DIR = Path(__file__).resolve(strict=True).parent.parent

chrome_options = webdriver.ChromeOptions() 
chrome_options.add_argument('--headless')
chrome_options.add_argument('--no-sandbox')
chrome_options.add_argument('--disable-dev-shm-usage')
chrome_options.add_argument("--remote-debugging-port=53906")
# chrome_options.add_argument("--remote-debugging-port=9591")
chrome_options.add_argument("--disable-gpu") 
chrome_options.add_argument('--ignore-certificate-errors')
chrome_options.add_argument('--allow-insecure-localhost')
chrome_options.add_argument('--allow-running-insecure-content')
driver = webdriver.Chrome('/usr/bin/chromedriver', options=chrome_options, service_args=["--verbose", "--log-path=/home/patoausa/app1/chromelog.txt"])
# driver = webdriver.Chrome('/usr/lib64/chromium-browser/chromedriver', options=chrome_options, service_args=["--verbose", "--log-path=/home/patoausa/app1/chromelog.txt"])
# driver = webdriver.Chrome('/usr/lib64/chromium-browser/chromedriver', service_args=["--verbose", "--log-path=/home/patoausa/app1/chromelog.txt"])

# PROJECT_ROOT = os.path.dirname(os.path.realpath('__file__'))
# DRIVER_BIN = os.path.join(PROJECT_ROOT, "chromedriver")
# chrome_options = webdriver.ChromeOptions()
# chrome_options.add_argument('--headless')
# chrome_options.add_experimental_option(
#     "excludeSwitches", ['enable-automation'])
# chrome_options.add_argument('--no-sandbox')
# chrome_options.add_argument("--disable-extensions")
# chrome_options.add_argument("--incognito")
# #driver = webdriver.Chrome(DRIVER_BIN, options=chrome_options)

MEDIA_ROOT = os.path.join(BASE_DIR, 'media')

@login_required
def dashboard(request):
    AB112 = ''
    patnof =''
    links=''
    claimlist =[]
    claims2=[]
    claim =[]
    claimcount=''
    pats= ''
    # if request.method == "POST" and Link.objects.count()>0:
    form1 = PatForm()
    form2 = PatentForm()
    if request.method == "POST":
            print(request.POST)
            # form1 = PatForm()
            # form2 = PatentForm()
            form1 = PatForm(request.POST or None)
            form2 = PatentForm(request.POST or None)
            if form1.is_valid() and form2.is_valid():
                A112 = form2.cleaned_data['A112']
                obj = form2.cleaned_data['obj']
                spec = form2.cleaned_data['spec']
                drw = form2.cleaned_data['drw']
                A102103 = form2.cleaned_data['A102103']
                print("AB112 is:\n", AB112)
                if A112 and obj and spec and drw and A102103:
                    print("All selected ")
                    AB112 = 'allselected'
                elif A112:
                    print("selected 112")
                    AB112 = "112"
                elif obj:
                    print("You selected obj")
                    AB112 = "obj"
                elif A102103:
                    print("You selected A102103")
                    AB112 = "A102103"
                elif spec:
                    print("You selected spec")
                    AB112 = "spec"
                elif drw:
                    print("You selected drw")
                    AB112 = "drw"
                else:
                    print("No selection")
                    AB112 = ''

                print("What chosen:", AB112)  

                patnof = form1.cleaned_data['patnof']
                print("Patent  ##:", patnof)
                driver = webdriver.Chrome(
                    executable_path=ChromeDriverManager().install(), chrome_options=chrome_options)
                driver.get('https://patents.google.com')
                xpath = '//input[@id="searchInput"]'
                searchBox = driver.find_element_by_xpath(xpath)
                searchBox.send_keys(patnof)
                searchBox.send_keys(Keys.ENTER)
                time.sleep(5)
                purl = driver.current_url
                print(f" Patent URL: {purl}\n")
                data = get(purl).text
                # data = response.text
                soup = BeautifulSoup(data, 'html.parser')
                claims = soup.findAll(id=True, class_="claim")
                #print("Claim from soup", claims)
                cln = len(claims)  # Total claims
                print("Total Claims:", cln)
                #print("First claim is :", claims[0])
                claimlist = []  # to hold array of stings of claims, an empty list is declared
                n = 0
                while n < cln:
                    # claimlist.append(claims[n].text)  # to genrate a list of strings
                    claimlist.append(claims[n].text.strip())  # to genrate a list of strings
                    # claimlist=claims[n].text
                    n += 1

                #print("path for docx file", patnopath)
                patnopath = os.path.join(BASE_DIR, 'patoa', 'static', 'media')
                claims2doc(claimlist, patnopath, AB112, patnof)

                patents=form1.save(commit = False )
                Patent.objects.create(patnof=patnof, claim_list=claimlist,  user=request.user)
                # patents.user=request.user
                # Patent.objects.create(patnof=patnof, claim_list=claimlist, user=user.username)
                # patents.save()
                pats = Patent.objects.values()
                patno = Patent.objects.values('patnof')
                patnof = patno.last()['patnof']
                print(" Value of patnof:", patnof)
                links = Patent.objects.values('claim_list').last()['claim_list']
                claimcount = str(len(links))
                print("Links  ::::", claimcount)
                # lastlink = links.last()
                # form2.save
            # form2 = PatForm(request.POST)
            # if form2.is_valid():
                # form.save(commit = True)   
                
                driver.quit()
            else:
                    print("Form is not valid")
            
    else:
                print("In the else block")
                form1 = PatForm()
                form2 = PatentForm()
        #     print("I am in else for Post Method")
    context = {
                'patnof': patnof,
                'links': links,
                #'lastlink': lastlink,
                'claimlist': claimlist,
                'claimcounter': claimcount,
                'form1' : form1,
                'form2' : form2,
                # 'pats' : pats,
            }
    # return render(request, 'patoa/result.html',context)
    return render(request, 'patoa/dashboard.html',context)
   


@login_required
def clear(request,username):
    # Patent.objects.all().delete()
    user =  User.objects.get(username=username)
    Patent.objects.filter(user=user).delete()
    return render(request, 'patoa/dashboard.html')


def claims2doc(claimslist, path, sel, patno):
    docu = Document()
    docu.styles['Normal'].font.size = Pt(12)
    docu.styles['Normal'].font.name = 'Arial'
    docu.styles['Heading 1'].font.color.rgb = RGBColor(0, 0, 0)

    def make_paragraph(text, size=10):
        p = docu.add_paragraph()
        p.add_run(text).font.size = Pt(size)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        #p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    def make_paragraph1(text, size=12):
        p = docu.add_paragraph()
        #p.add_run(text).font.size = Pt(size)
        p.add_run(text).bold = True
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    def make_paragraphL(text, size=12):
        p = docu.add_paragraph()
        #p.add_run(text).font.size = Pt(size)
        p.add_run(text).bold = True
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if sel == 'allselected':
        print(sel)
        make_paragraph1('Specification')
        make_paragraph('The title of the invention is not descriptive.  A new title is required that is clearly indicative of the invention' + 
        'to which the claims are directed. Suggested title “a manipulator used to drive a surgical device that treats a body tissue ”.\n', size=12)

        make_paragraph1('Drawing')
        make_paragraph('The drawings are objected to because Fig. 2-14 are not showing the labels and/legends in the picture clearly and the pictures are hazy and vague.'+  
        'Corrected drawing sheets in compliance with 37 CFR 1.121(d) are required in reply to the Office action to avoid abandonment of the application.'+  
        'Any amended replacement drawing sheet should include all of the figures appearing on the immediate prior version of the sheet, even if only one ' + 
        'figure is being amended. The figure or figure number of an amended drawing should not be labeled as “amended.” If a drawing figure is to be canceled,'+ 
        'the appropriate figure must be removed from the replacement sheet, and where necessary, the remaining figures must be renumbered and appropriate changes,' + 
        'made to the brief description of the several views of the drawings for consistency. Additional replacement sheets may be necessary to show the renumbering'+ 
        'of the remaining figures. Each drawing sheet submitted after the filing date of an application must be labeled in the top margin as either “Replacement Sheet” or “New Sheet” pursuant to 37 CFR 1.121(d). If the changes are not accepted by the examiner, the applicant will be notified and informed of any required corrective action in the next Office action. The objection to the drawings will not be held in abeyance.\n', size=12)


        #docu.add_heading("Claim rejection under 35 USC 112")
        make_paragraph1('Claim rejection under 35 USC 112')
        #para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        make_paragraph1(
            'The following is a quotation of the first paragraph of 35 U.S.C. 112(a):',size=12)
        make_paragraph('(a) IN GENERAL.—The specification shall contain a written description of the invention, and of the manner ' +
                       'and process of making and using it, in such full, clear, concise, and exact terms as to enable any person skilled'+
                        'in the art to which it pertains,' +
                       'or with which it is most nearly connected, to make and use the same, and shall set forth the best mode contemplated'+
                        'by the inventor or joint inventor of carrying out the invention.', size=10)
        make_paragraph1('The following is a quotation of 35 U.S.C. 112(b):',size=12)
        make_paragraph('(b)  CONCLUSION.—The specification shall conclude with one or more claims particularly pointing out and distinctly'+
         'claiming the subject matter which the inventor or a joint inventor regards as the invention.\n'+
         'The following is a quotation of 35 U.S.C. 112 (pre-AIA), second paragraph:'+
        'The specification shall conclude with one or more claims particularly pointing out and distinctly '+
        'claiming the subject matter which the applicant regards as his invention.', size=10)

        make_paragraph(
            'Claims 1-19   are rejected under 35 U.S.C. 112(b) or 35 U.S.C. 112 (pre-AIA),'+
            'second paragraph, as being indefinite for failing to particularly point out and '+
            'distinctly claim the subject matter which the inventor or a joint inventor, or '+
            'for pre-AIA the applicant regards as the invention. \n', size=12)


        make_paragraph1('Claim rejection under 35 USC 102',size=12)
        #para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        make_paragraph('The following is a quotation of the appropriate paragraphs of 35 U.S.C. 102 that form the basis for the rejections under this section made in this Office action:')
        make_paragraph('A person shall be entitled to a patent unless –' +
                       '(a)(1) the claimed invention was patented, described in a printed publication, or' +
                       'in public use, on sale or otherwise available to the public before the effective' + 
                       'filing date of the claimed invention.', size=8
                       )
        make_paragraph1('Claims 1-19 are rejected under 35 U.S.C. 102(a)(1) as being anticipated by  XXXXX et al (US )\n', size=12)

        para = docu.add_paragraph('Claim rejection under 35 USC 103')
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        make_paragraph1('The following is a quotation of 35 U.S.C. 103 which forms the basis for all obviousness rejections set forth in this Office action:')
        make_paragraph('A patent for a claimed invention may not be obtained, notwithstanding that the claimed invention is not identically disclosed as set forth in section 102 of this title'+
        'if the differences between the claimed invention and the prior art are such that the claimed invention as a whole would have been obvious before the effective filing date of the claimed'+
        'invention to a person having ordinary skill in the art to which the claimed invention pertains.'+
        'Patentability shall not be negated by the manner in which the invention was made.\n', size=12)

        make_paragraph('Claims 1-11 are rejected under 35 U.S.C. 103 as being unpatentable over XXXXXXX (US 20160142003) in view of XXXXXXX. (US ).\n', size=12)

    elif sel == '112':
        # para = docu.add_paragraph('Claim rejection under 35 USC 112')
        make_paragraph1('Claim rejection under 35 USC 112')
        #para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        make_paragraph1(
            'The following is a quotation of the first paragraph of 35 U.S.C. 112(a):')
        make_paragraph('(a) IN GENERAL.—The specification shall contain a written description of the invention, and of the manner ' +
                       'and process of making and using it, in such full, clear, concise, and exact terms as to enable any person skilled in the art to which it pertains,' +
                       'or with which it is most nearly connected, to make and use the same,'+ 
                       'and shall set forth the best mode contemplated by the inventor or joint inventor of carrying out the invention. \n \n', size=10)
        make_paragraph1('The following is a quotation of 35 U.S.C. 112(b): \n')
        make_paragraph('(b)  CONCLUSION.—The specification shall conclude with one or more claims particularly pointing out and distinctly \n'+
         'claiming the subject matter which the inventor or a joint inventor regards as the invention.\n'+
         'The following is a quotation of 35 U.S.C. 112 (pre-AIA), second paragraph:'+
        'The specification shall conclude with one or more claims particularly pointing out and distinctly claiming the subject matter which the applicant regards as his invention.\n', size=10)

        make_paragraph(
            'Claims 1-19   are rejected under 35 U.S.C. 112(b) or 35 U.S.C. 112 (pre-AIA),'+
            'second paragraph, as being indefinite for failing to particularly point out and '+
            'distinctly claim the subject matter which the inventor or a joint inventor, \n', size=12)
    elif sel == 'A102103':
        make_paragraph1('Claim rejection under 35 USC 102')
        # para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        make_paragraph('The following is a quotation of the appropriate paragraphs of 35 U.S.C. 102 that form the basis for the rejections under this section made in this Office action:')
        make_paragraph('A person shall be entitled to a patent unless –' +
                       '(a)(1) the claimed invention was patented, described in a printed publication, '+ 
                       'or in public use, on sale or otherwise available to the public before the '+ 
                       'effective filing date of the claimed invention.\n', size=10)
        make_paragraph1('Claims 1-19 are rejected under 35 U.S.C. 102(a)(1) as being anticipated by  XXXXX et al (US )\n', size=12)

        make_paragraph1('Claim rejection under 35 USC 103')
        # para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        make_paragraph1('The following is a quotation of 35 U.S.C. 103 which forms the basis for all obviousness rejections set forth in this Office action:')
        make_paragraph('A patent for a claimed invention may not be obtained, notwithstanding that the claimed invention is not identically disclosed as set forth in section 102 of this title'+
        'if the differences between the claimed invention and the prior art are such that the claimed invention as a whole would have been obvious before the effective filing date of the claimed'+
        'invention to a person having ordinary skill in the art to which the claimed invention pertains.'+
        'Patentability shall not be negated by the manner in which the invention was made.\n', size=10)

        make_paragraph1('Claims 1-11 are rejected under 35 U.S.C. 103 as being unpatentable over XXXXXXX (US 20160142003) in view of XXXXXXX. (US ).\n', size=12)

    elif sel == 'spec':
        make_paragraph1('Specification')
        make_paragraph('The title of the invention is not descriptive.  A new title is required that is clearly indicative of the invention' + 
        'to which the claims are directed. Suggested title “a manipulator used to drive a surgical device that treats a body tissue ”.\n', size=12)

        print(sel)
    elif sel == 'drw':
        make_paragraph1('Drawing')
        make_paragraph('The drawings are objected to because Fig. 2-14 are not showing the labels and/legends in the picture clearly and the pictures are hazy and vague.'+  
        'Corrected drawing sheets in compliance with 37 CFR 1.121(d) are required in reply to the Office action to avoid abandonment of the application.'+  
        'Any amended replacement drawing sheet should include all of the figures appearing on the immediate prior version of the sheet, even if only one ' + 
        'figure is being amended. The figure or figure number of an amended drawing should not be labeled as “amended.” If a drawing figure is to be canceled,'+ 
        'the appropriate figure must be removed from the replacement sheet, and where necessary, the remaining figures must be renumbered and appropriate changes,' + 
        'made to the brief description of the several views of the drawings for consistency. Additional replacement sheets may be necessary to show the renumbering'+ 
        'of the remaining figures. Each drawing sheet submitted after the filing date of an application must be labeled in the top margin as either “Replacement Sheet” or “New Sheet” pursuant to 37 CFR 1.121(d). If the changes are not accepted by the examiner, the applicant will be notified and informed of any required corrective action in the next Office action. The objection to the drawings will not be held in abeyance.\n', size=12)

    elif sel == 'obj':
        make_paragraph1('Objection')
        make_paragraph(' Minor informalites.\n', size=12)

    else:
        print("None")
        AB112 = ''

    p = docu.add_paragraph()
    for i in range(len(claimslist)):
        claim = claimslist[i].strip()
        for j in range(len(claim)):
            if claim[j] == ".":
                claim = claim[j:]
                break
        p.add_run(f"Regarding claim {i+1}").bold = True
        p.add_run(f"{claim} \n")
        p.add_run(f" \n")
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        #p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    all_paras = docu.paragraphs
    #print(f"Added all claims. No of Paras:", len(all_paras))
    filename = ("/" + str(patno) + '.docx')
    docu.save(path+filename)

@login_required
def profile(request, username):
    # user =  Patent.objects.all().values('id')[0]['id']
    user =  User.objects.get(username=username)
    patents = Patent.objects.filter(user=user)
    if patents.exists():
    # patno =   Patent.objects.values('patnof')[0]['patnof']
        claims = patents.values('claim_list')[0]['claim_list']
        clms = patents.values('claim_list')
        # claims=clms['claim_list']
        # patnof = patents.values('patnof').last()['patnof']
        return render(request,'patoa/profile.html',{'username':username,
                    'patents': patents, 'claims':claims})
    else:
        
        return render(request,'patoa/dashboard.html')

# @login_required
def patclaim(request, patnof):
    pf = Patent.objects.all().filter(patnof=patnof)
    clm=pf.values('claim_list')[0]['claim_list']
    patno =  patnof
    return render(request,'patoa/patclaim.html',{'claims':clm, 'patnof':patno})

@login_required
def remov(request,patnof):
    Patent.objects.filter(patnof=patnof).delete()
    return render(request, 'patoa/profile.html')
    # return redirect('')